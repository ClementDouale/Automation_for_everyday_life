import os
from bs4 import BeautifulSoup
from docx import Document
import requests
from zipfile import ZipFile

# Function to download and add image to the document
def add_image_to_doc(doc, img_url):
    # Code to download image and add to the document
    ## pass
     # Check if the image URL is absolute or relative
    if not img_url.startswith(('http://', 'https://')):
        img_url = base_url + img_url  # Adjust this based on your HTML structure

    # Download the image
    try:
        response = requests.get(img_url)
        response.raise_for_status()  # Raise an error for bad status codes
        image_bytes = BytesIO(response.content)

        # Add the image to the document
        doc.add_picture(image_bytes)
    except requests.RequestException as e:
        print(f"Error downloading image {img_url}: {e}")

# Extract the HTML files from the zip
with ZipFile('path_to_your_zip_file.zip', 'r') as zip_ref:
    zip_ref.extractall('extracted_folder')

# Create a new Word document
doc = Document()

# Loop through each HTML file and add its content to the Word document
for filename in os.listdir('extracted_folder'):
    if filename.endswith('.html'):
        with open(os.path.join('extracted_folder', filename), 'r', encoding='utf-8') as file:
            soup = BeautifulSoup(file, 'html.parser')
            # Extract and add text to the document
            text = soup.get_text()
            doc.add_paragraph(text)
            # Find and handle images
            for img_tag in soup.find_all('img'):
                img_url = img_tag['src']
                add_image_to_doc(doc, img_url)

# Save the document
doc.save('output.docx')
